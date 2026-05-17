"""File parsing utilities for document uploads."""

from typing import Optional
from fastapi import UploadFile, HTTPException
from docx import Document
import io


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text from a .docx file.
    
    Args:
        file: Uploaded .docx file
        
    Returns:
        Extracted text content
        
    Raises:
        HTTPException: If file is invalid or cannot be parsed
    """
    try:
        # Read file content
        contents = await file.read()
        
        # Load document from bytes
        doc = Document(io.BytesIO(contents))
        
        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        # Join all text
        full_text = "\n".join(paragraphs)
        
        if not full_text.strip():
            raise HTTPException(status_code=400, detail="Document contains no extractable text")
        
        return full_text
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse .docx file: {str(e)}")
    finally:
        await file.seek(0)


async def validate_and_extract_text(file: UploadFile) -> tuple[str, str]:
    """
    Validate file type and extract text.
    
    Args:
        file: Uploaded file
        
    Returns:
        Tuple of (extracted_text, original_filename)
        
    Raises:
        HTTPException: If file type is unsupported
    """
    # Check file extension
    filename = file.filename or "unknown"
    extension = filename.lower().split('.')[-1] if '.' in filename else ''
    
    if extension not in ['docx']:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file format: .{extension}. Currently only .docx is supported."
        )
    
    # Extract text based on file type
    if extension == 'docx':
        text = await extract_text_from_docx(file)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")
    
    return text, filename
